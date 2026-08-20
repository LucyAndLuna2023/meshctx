#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""v4 报告 md → docx（python-docx 程序化，符合 b2/b3 铁律）

用法: python3 md2docx_report.py <in.md> <out.docx>
"""
import sys, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main(src, dst):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Microsoft YaHei'
    st.font.size = Pt(10)
    for sec in doc.sections:
        sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)
        sec.top_margin = Cm(1.0); sec.bottom_margin = Cm(1.0)
    lines = open(src, encoding='utf-8').read().splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.strip():
            i += 1; continue
        # 表格块
        if ln.lstrip().startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s:\-|]+\|\s*$', lines[i+1]):
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(cells); i += 1
            if rows:
                t = doc.add_table(rows=len(rows)-1, cols=len(rows[0]))
                t.style = 'Table Grid'
                for r_idx, row in enumerate(rows[1:]):
                    for c_idx in range(min(len(row), len(rows[0]))):
                        cell = t.rows[r_idx].cells[c_idx]
                        cell.text = row[c_idx]
                        for p in cell.paragraphs:
                            p.paragraph_format.space_after = Pt(2)
                doc.add_paragraph('')
            continue
        # 标题
        m = re.match(r'^(#{1,4})\s+(.*)$', ln)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(m.group(2), level=min(level, 4))
            if level == 1:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1; continue
        # 引用
        if ln.startswith('>'):
            p = doc.add_paragraph(ln.lstrip('> ').strip())
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(3)
            i += 1; continue
        # 分隔线
        if re.match(r'^\s*[-*_]{3,}\s*$', ln):
            i += 1; continue
        # 列表
        m = re.match(r'^(\s*)[-*+]\s+(.*)$', ln)
        if m:
            p = doc.add_paragraph(m.group(2), style='List Bullet')
            i += 1; continue
        m = re.match(r'^(\s*)\d+\.\s+(.*)$', ln)
        if m:
            p = doc.add_paragraph(m.group(2), style='List Number')
            i += 1; continue
        # 普通段落（去掉 markdown 加粗符号）
        txt = re.sub(r'\*\*([^*]+)\*\*', r'\1', ln)
        txt = re.sub(r'`([^`]+)`', r'\1', txt)
        doc.add_paragraph(txt)
        i += 1
    doc.save(dst)
    print(f'OK: {dst} ({len(lines)} lines)')

if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2])
