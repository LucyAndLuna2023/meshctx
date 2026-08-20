"""
meshctx Chat 工具引擎 v2 — 原生 OpenAI function calling 格式
对标 Hermes: read_file, write_file, search_files, terminal, web_search, list_dir
"""
import logging
import os, re, json, shlex, subprocess, urllib.request, urllib.parse
from pathlib import Path
from typing import Dict, Optional, List

logger = logging.getLogger("meshctx.chat_tools")


def _detect_desktop() -> str:
    """检测用户真实桌面目录（跨平台）。

    Linux: XDG user-dirs.dirs 优先（如 ~/桌面）→ ~/Desktop / ~/桌面
    Windows: 注册表 User Shell Folders（OneDrive 重定向）→ ~/OneDrive/Desktop → ~/Desktop / 中文桌面
    macOS: ~/Desktop
    """
    try:
        home = Path.home()
        if os.name == "nt":
            # Windows: OneDrive 等会把桌面重定向到注册表指定目录
            try:
                import winreg
                with winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                                    r"Software\Microsoft\Windows\CurrentVersion\Explorer\User Shell Folders") as k:
                    val, _ = winreg.QueryValueEx(k, "Desktop")
                val = val.replace("%USERPROFILE%", str(home))
                cand = Path(val)
                if cand.is_dir():
                    return str(cand)
            except Exception:
                pass
            for cand in (home / "OneDrive" / "Desktop", home / "Desktop",
                         home / "OneDrive" / "桌面", home / "桌面"):
                if cand.is_dir():
                    return str(cand)
            return str(home / "Desktop")
        # Linux: XDG 本地化目录（中文环境为 ~/桌面）
        cfg = home / ".config" / "user-dirs.dirs"
        if cfg.exists():
            for line in cfg.read_text(encoding="utf-8", errors="replace").splitlines():
                line = line.strip()
                if line.startswith("XDG_DESKTOP_DIR="):
                    val = line.split("=", 1)[1].strip().strip('"').strip("'")
                    if val.startswith("$HOME"):
                        cand = home / val.replace("$HOME", "").lstrip("/")
                    else:
                        cand = Path(val)
                    if cand.is_dir():
                        return str(cand)
        # macOS / Linux 兜底
        for cand in (home / "Desktop", home / "桌面"):
            if cand.is_dir():
                return str(cand)
    except Exception:
        pass
    return str(Path.home() / "Desktop")


USER_DESKTOP = _detect_desktop()


def _normalize_desktop_path(path: str) -> str:
    """若真实桌面为本地化目录(如 ~/桌面)而模型误用 ~/Desktop（或反向），自动改写到真实桌面"""
    try:
        p = Path(path).expanduser()
        if not p.is_absolute():
            return path
        home = Path.home()
        real = Path(USER_DESKTOP)
        for cand_name in ("Desktop", "桌面", "OneDrive/Desktop"):
            cand = home / cand_name
            if cand == real:
                continue
            try:
                rel = p.relative_to(cand)
                return str(real / rel)
            except ValueError:
                continue
    except Exception:
        pass
    return path

_DSML_INVOKE_RE = re.compile(r'<invoke\s+name="([^"]+)"[^>]*>(.*?)</invoke>', re.S)
_DSML_PARAM_RE = re.compile(r'<parameter\s+name="([^"]+)"[^>]*>(.*?)</parameter>', re.S)
_DSML_BLOCK_RE = re.compile(r'<tool_calls>.*?</tool_calls>', re.S)


_DSML_TAG_NOISE_RE = re.compile(r'<(\s*/?\s*)(?:DSML[｜|]*)?[｜|]+([A-Za-z_][A-Za-z_0-9]*)')


def _clean_dsml_tags(text: str) -> str:
    """清理 deepseek 在 DSML 标签名中注入的全角竖线 ｜(U+FF5C)，如 <｜invoke> → <invoke"""
    try:
        return _DSML_TAG_NOISE_RE.sub(lambda m: "<" + m.group(1) + m.group(2), text or "")
    except Exception:
        return text


def parse_dsml_tool_calls(text: str, prefix: str = "dsml") -> list:
    """解析 Hermes/DeepSeek 风格的 DSML <tool_calls><invoke name=..> 文本块为工具调用列表。

    返回 [{id, name, arguments(dict)}]。当模型不返回原生 function-call、
    而是把工具调用写成 XML 文本（含被注入全角 ｜ 的变体）时，CLI/UI 共用此解析保证工具仍会执行。
    """
    calls = []
    try:
        text = _clean_dsml_tags(text or "")
        for i, m in enumerate(_DSML_INVOKE_RE.finditer(text)):
            name = m.group(1).strip()
            body = m.group(2)
            args = {}
            for pm in _DSML_PARAM_RE.finditer(body):
                k = pm.group(1).strip()
                v = pm.group(2).strip()
                try:
                    args[k] = json.loads(v)
                except Exception:
                    args[k] = v
            if name:
                calls.append({"id": f"{prefix}_{i}", "name": name, "arguments": args})
    except Exception:
        return []
    return calls


def strip_dsml_tool_calls(text: str) -> str:
    """移除文本中的 <tool_calls>...</tool_calls> 块（保留块外内容）"""
    if not text:
        return text
    try:
        text = _clean_dsml_tags(text)
        if "<tool_calls>" not in text:
            return text
        return _DSML_BLOCK_RE.sub("", text).strip()
    except Exception:
        return text


# ═══════════════════════════════════════════════════
# 工具执行函数
# ═══════════════════════════════════════════════════

def _read_file(path: str, limit: int = 100) -> str:
    try:
        p = Path(path).expanduser()
        if not p.exists():
            return f"文件不存在: {path}"
        with open(p, "r", encoding="utf-8", errors="replace") as f:
            lines = f.readlines()
        if len(lines) > limit:
            return "".join(lines[:limit]) + f"\n... (共{len(lines)}行, 显示前{limit}行)"
        return "".join(lines)
    except Exception as e:
        return f"读取失败: {e}"

def _write_file(path: str, content: str, if_exists: str = "rename") -> str:
    try:
        p = Path(_normalize_desktop_path(path)).expanduser()
        p.parent.mkdir(parents=True, exist_ok=True)
        renamed = False
        if if_exists == "rename" and p.exists():
            stem, ext = p.stem, p.suffix
            n = 1
            while p.with_name(f"{stem}({n}){ext}").exists():
                n += 1
            p = p.with_name(f"{stem}({n}){ext}")
            renamed = True
        p.write_text(content, encoding="utf-8")
        suffix = " (文件已存在，另存为新文件)" if renamed else ""
        return f"已写入: {p} ({len(content)}字符){suffix}"
    except Exception as e:
        return f"写入失败: {e}"

def _save_docx(path: str, title: str, content: str) -> str:
    """直接生成 Word(.docx) 报告（服务器端 python-docx）— 重名自动另存"""
    import os as _os
    try:
        from docx import Document
    except ImportError:
        return "生成Word失败: 服务器未安装 python-docx，请改用 terminal 安装或使用 write_file 输出文本报告"
    try:
        p = Path(_normalize_desktop_path(path)).expanduser()
        p.parent.mkdir(parents=True, exist_ok=True)
        renamed = False
        if p.exists():
            stem, ext = p.stem, p.suffix
            n = 1
            while p.with_name(f"{stem}({n}){ext}").exists():
                n += 1
            p = p.with_name(f"{stem}({n}){ext}")
            renamed = True
        doc = Document()
        if title:
            doc.add_heading(str(title), level=0)
        for line in (content or "").splitlines():
            line = line.rstrip()
            if not line.strip():
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif line.startswith("* "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line.strip())
        doc.save(str(p))
        size = _os.path.getsize(str(p))
        suffix = " (文件已存在，另存为新文件)" if renamed else ""
        return f"已生成Word报告: {p} ({size}字节){suffix}"
    except Exception as e:
        return f"生成Word失败: {e}"


def _list_dir(path: str) -> str:
    try:
        p = Path(path).expanduser()
        if not p.exists():
            return f"目录不存在: {path}"
        items = sorted(p.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
        lines = []
        for item in items[:50]:
            t = "📁" if item.is_dir() else "📄"
            size = ""
            if item.is_file():
                s = item.stat().st_size
                size = f" ({s:,}B)" if s < 1024 else f" ({s/1024:.1f}KB)"
            lines.append(f"  {t} {item.name}{size}")
        return "\n".join(lines) if lines else "(空目录)"
    except Exception as e:
        return f"列出失败: {e}"

def _default_path_dirs() -> list:
    """跨平台系统命令目录：Linux/macOS 的 POSIX bin 目录；Windows 返回空（用系统 PATH 即可）"""
    return [d for d in ("/usr/local/sbin", "/usr/local/bin", "/usr/sbin", "/usr/bin", "/sbin", "/bin")
            if os.path.isdir(d)]


def _run_cmd(cmd: str) -> str:
    try:
        env = os.environ.copy()
        _extra = _default_path_dirs()
        if _extra:
            env['PATH'] = os.pathsep.join(_extra) + os.pathsep + env.get('PATH', '')
        # 全盘扫描拦截（机制层，不依赖模型自觉）：find/grep/du/locate 以根目录 / 为目标必然超时
        _full_scan = re.search(r'\b(?:find|grep|du|locate)(?:\s+[^;&|]*)?\s+/(?:\s|$)', cmd or '')
        if _full_scan:
            return ("[安全拦截] 已阻止从根目录 '/' 的全盘扫描（30秒必超时且无意义）。"
                    "请改用精确路径，例如: ls ~/.hermes/profiles/quant、search_files 工具、"
                    "find ~/.hermes -name '...'")
        # 危险命令检测（CLI/UI 统一，与 /api/terminal 一致）；shell=True 支持 &&/管道/重定向
        from src.core.sandbox import CodeScanner
        ok, err = CodeScanner.scan_bash(cmd)
        if not ok:
            return f"[安全拦截] 危险命令已被拦截: {err}"
        r = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=30, cwd=os.getcwd(), env=env)
        out = r.stdout[:3000]
        if r.stderr:
            out += "\n[stderr]\n" + r.stderr[:500]
        return out or f"(exit={r.returncode})"
    except subprocess.TimeoutExpired:
        return ("命令超时(30秒)。请缩小范围重试，例如限制目录深度、指定具体路径，"
                "不要对整个文件系统做遍历。")
    except Exception as e:
        return f"执行失败: {e}"

def _search_files(pattern: str, path: str = ".", glob: str = "*") -> str:
    try:
        p = Path(path).expanduser()
        if not p.exists():
            return f"目录不存在: {path}"
        matches = []
        for f in p.rglob(glob):
            if f.is_file() and f.stat().st_size < 1024*1024:
                try:
                    content = f.read_text(errors="replace")
                    if pattern.lower() in content.lower():
                        matches.append(str(f))
                except Exception:
                    logger.debug("chat_tools error", exc_info=True)
                    pass
        if not matches:
            return f"未找到包含 '{pattern}' 的文件"
        return "\n".join(matches[:20])
    except Exception as e:
        return f"搜索失败: {e}"

def _get_proxy() -> str:
    """读取搜索代理配置: meshctx.yaml > env > None"""
    import yaml
    config_paths = [
        Path(__file__).parent.parent / "meshctx.yaml",
        Path.home() / ".meshctx" / "config.yaml",
    ]
    for cp in config_paths:
        if cp.exists():
            try:
                cfg = yaml.safe_load(cp.read_text())
                proxy = cfg.get("search", {}).get("proxy", "")
                if proxy:
                    return proxy
            except Exception:
                pass
    # 环境变量回退
    for env_var in ["ALL_PROXY", "HTTPS_PROXY", "HTTP_PROXY"]:
        val = os.environ.get(env_var, "")
        if val:
            return val
    return ""


def _web_search(query: str) -> str:
    """多引擎网页搜索: ddgs(代理翻墙) → cn.bing.com(直连兜底)"""
    date_keywords = ['今天', '今日', '最新', '实时', '当前', '目前']
    if any(k in query for k in date_keywords):
        from datetime import datetime
        today = datetime.now().strftime('%Y年%m月%d日')
        if today not in query:
            query = f"{query} {today}"

    proxy = _get_proxy()

    raw_results = ""

    # 引擎1: ddgs 库 (DuckDuckGo API, 质量最高, 需代理翻墙)
    if proxy:
        try:
            result = _search_ddgs_api(query, proxy)
            if result:
                raw_results = result
        except Exception:
            logger.debug("chat_tools ddgs error", exc_info=True)

    # 引擎2: cn.bing.com 直连 (中文/兜底)
    if not raw_results:
        try:
            result = _search_bing_cn(query)
            if result:
                raw_results = result
        except Exception:
            logger.debug("chat_tools bing cn error", exc_info=True)

    if not raw_results:
        return "搜索失败: 所有搜索引擎均不可用"

    # ── P1-4: URL 去重 + 引用溯源 ──
    # 结果行格式: "N. 标题\n   snippet\n   url"，按 url 去重，保留首个出现
    seen: set = set()
    dedup_lines: list = []
    idx = 0
    for line in raw_results.split("\n"):
        line = line.rstrip()
        if line.strip().startswith("http"):
            url = line.strip()
            if url in seen:
                continue  # 重复 URL 丢弃
            seen.add(url)
            idx += 1
            dedup_lines.append(f"[{idx}] {url}")
        else:
            dedup_lines.append(line)
    # 重新编号标题行（去重后编号可能不连续，统一重排：标题行→去掉旧编号前缀）
    out_lines = []
    for line in dedup_lines:
        m = re.match(r"^\d+\. (.*)$", line)
        if m:
            out_lines.append(line)  # 保留原编号，引用溯源以 [N] 链接为准
        else:
            out_lines.append(line)
    return "\n".join(out_lines)


def _search_ddgs_api(query: str, proxy: str) -> str:
    """DuckDuckGo API — 通过 ddgs 库, 质量最高"""
    # 环境变量仅在函数内临时设置（finally 恢复），避免污染全局 env（004 隐患#6 同款问题）
    import os as _os
    _saved = {k: _os.environ.get(k) for k in ("ALL_PROXY", "HTTPS_PROXY", "HTTP_PROXY")}
    for _k in ("ALL_PROXY", "HTTPS_PROXY", "HTTP_PROXY"):
        _os.environ[_k] = proxy
    try:
        try:
            from ddgs import DDGS
            with DDGS(timeout=15) as ddgs:
                results = list(ddgs.text(query, max_results=6))
            if not results:
                return ""
            lines = []
            for i, r in enumerate(results):
                title = r.get("title", "")
                body = r.get("body", "")[:200]
                href = r.get("href", "")
                if title:
                    lines.append(f"{i+1}. {title}\n   {body}\n   {href}")
            return "\n".join(lines) if lines else ""
        except ImportError:
            return ""
        except Exception as e:
            return ""
    finally:
        # 恢复环境变量，避免污染全局
        for _k, _v in _saved.items():
            if _v is None:
                _os.environ.pop(_k, None)
            else:
                _os.environ[_k] = _v



def _search_bing_cn(query: str) -> str:
    """cn.bing.com 直连 — 中文内容"""
    url = f"https://cn.bing.com/search?q={urllib.parse.quote(query)}"
    req = urllib.request.Request(url, headers={
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    })
    with urllib.request.urlopen(req, timeout=10) as resp:
        raw = resp.read().decode(errors='replace')
    return _parse_bing_html(raw)


def _parse_bing_html(raw: str) -> str:
    """解析 Bing HTML 搜索结果"""
    import html as _html
    html_text = _html.unescape(raw)
    algos = re.findall(r'<li class="b_algo"[^>]*>(.*?)</li>', html_text, re.DOTALL)
    if not algos:
        return ""
    results = []
    for algo in algos[:8]:
        title_m = re.search(r'<h2[^>]*><a[^>]*>(.*?)</a>', algo, re.DOTALL)
        title = re.sub(r'<[^>]+>', '', title_m.group(1)).strip() if title_m else ""
        url_m = re.search(r'<a[^>]*href="(https?://[^"]+)"', algo)
        url_text = url_m.group(1) if url_m else ""
        cap_m = re.search(r'<div class="b_caption"[^>]*>(.*?)</div>', algo, re.DOTALL)
        snippet = ""
        if cap_m:
            raw_text = re.sub(r'<[^>]+>', ' ', cap_m.group(1))
            snippet = re.sub(r'\\s+', ' ', raw_text).strip()[:250]
        if title and snippet:
            results.append(f"{len(results)+1}. {title}\n   {snippet}\n   {url_text}")
        elif title:
            results.append(f"{len(results)+1}. {title}\n   {url_text}")
    return "\n".join(results[:6]) if results else ""

# ═══════════════════════════════════════════════════
# OpenAI Function Calling 工具定义
# ═══════════════════════════════════════════════════

TOOLS_OPENAI = [
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "读取文件内容。用于查看代码、配置、日志等文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "文件路径（绝对路径或相对路径）"},
                    "limit": {"type": "integer", "description": "最多显示行数，默认100", "default": 100}
                },
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": "写入文件。会自动创建父目录。生成报告/文档时若目标文件已存在且不想覆盖，传 if_exists='rename' 自动另存为新文件(文件名加 (1) 后缀)。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "文件路径"},
                    "content": {"type": "string", "description": "要写入的内容"},
                    "if_exists": {"type": "string", "description": "目标已存在时的处理方式: overwrite=覆盖, rename=自动另存新文件(默认)", "enum": ["overwrite", "rename"], "default": "rename"}
                },
                "required": ["path", "content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_dir",
            "description": "列出目录中的文件和子目录。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "目录路径，默认当前目录", "default": "."}
                },
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "run_cmd",
            "description": "执行终端命令。用于编译、安装包、git操作等。命令在30秒超时内运行。",
            "parameters": {
                "type": "object",
                "properties": {
                    "cmd": {"type": "string", "description": "要执行的shell命令"}
                },
                "required": ["cmd"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_files",
            "description": "在文件中搜索指定内容。用于查找函数定义、错误信息等。",
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {"type": "string", "description": "搜索文本（不区分大小写）"},
                    "path": {"type": "string", "description": "搜索起始目录，默认当前目录", "default": "."},
                    "glob": {"type": "string", "description": "文件名过滤，如 *.py 只搜索Python文件", "default": "*"}
                },
                "required": ["pattern"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "联网搜索网页信息。查股票、新闻、天气、文档等实时数据时优先使用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "搜索关键词"}
                },
                "required": ["query"]
            }
        }
    },
]

# 工具执行器映射
TOOL_EXECUTORS = {
    "read_file": lambda args: _read_file(args.get("path", ""), args.get("limit", 100)),
    "write_file": lambda args: _write_file(args.get("path", ""), args.get("content", ""), args.get("if_exists", "rename")),
    "list_dir": lambda args: _list_dir(args.get("path", ".")),
    "run_cmd": lambda args: _run_cmd(args.get("cmd", "")),
    "terminal": lambda args: _terminal(args.get("cmd", ""), args.get("workdir", ""), args.get("timeout", 60)),
    "search_files": lambda args: _search_files(args.get("pattern", ""), args.get("path", "."), args.get("glob", "*")),
    "web_search": lambda args: _web_search(args.get("query", "")),
    "web_extract": lambda args: _web_extract(args.get("url", "")),
    "remote_read": lambda args: _remote_read(args.get("path", ""), args.get("host", ""), args.get("user", ""), args.get("password", ""), args.get("port", 22)),
    "remote_write": lambda args: _remote_write(args.get("path", ""), args.get("content", ""), args.get("host", ""), args.get("user", ""), args.get("password", ""), args.get("port", 22)),
    "remote_exec": lambda args: _remote_exec(args.get("cmd", ""), args.get("host", ""), args.get("user", ""), args.get("password", ""), args.get("port", 22)),
    "browser_navigate": lambda args: _browser_navigate(args.get("url", ""), _BROWSER_CACHE),
    "browser_snapshot": lambda args: _browser_snapshot(_BROWSER_CACHE),
    "browser_click": lambda args: "[浏览器未授权] 请先在 UI 授权浏览器控制",
    "browser_type": lambda args: "[浏览器未授权] 请先在 UI 授权浏览器控制",
    "save_memory": lambda args: _save_memory(args.get("fact", "")),
    "save_docx": lambda args: _save_docx(args.get("path", ""), args.get("title", ""), args.get("content", "")),
}

TOOL_ICONS = {
    "read_file": "📖", "write_file": "📝", "list_dir": "📂",
    "run_cmd": "⚡", "terminal": "⚡", "search_files": "🔍", "web_search": "🌐", "save_docx": "📄",
    "web_extract": "🌐", "remote_read": "🖥️", "remote_write": "🖥️", "remote_exec": "🖥️",
    "browser_navigate": "🧭", "browser_snapshot": "🧭", "browser_click": "🧭", "browser_type": "🧭",
    "save_memory": "🧠",
}

def execute_tool(tool_name: str, arguments: Dict) -> str:
    """执行工具并返回结果"""
    fn = TOOL_EXECUTORS.get(tool_name)
    if not fn:
        return f"未知工具: {tool_name}"
    try:
        return fn(arguments)
    except Exception as e:
        return f"工具执行错误: {e}"

_TOOLS_PROMPT_TEMPLATE = """可用工具: read_file, write_file, save_docx, list_dir, run_cmd, terminal, search_files, web_search, web_extract, remote_read, remote_write, remote_exec, browser_navigate, browser_snapshot, save_memory
生成 Word 报告类文件时直接用 save_docx（服务器端生成 .docx，不需要写 python 脚本）；重名自动另存。
你运行在用户本地机器上，有完整的文件系统访问权限。
用户桌面目录: __USER_DESKTOP__（生成报告/文档默认保存到这里）
看到需要读文件/执行命令/搜索代码时，直接用工具！不要说"我无法访问"。
回复简洁有用，中文优先。
write_file 默认自动处理重名：目标文件已存在时自动另存为新文件(文件名加 (1) 后缀)，绝不覆盖旧文件；只有明确修改已有文件时才传 if_exists='overwrite'。"""


def get_tools_prompt() -> str:
    """生成简洁的工具列表文本（用于系统提示中的快速参考）— 桌面路径动态注入 USER_DESKTOP"""
    return _TOOLS_PROMPT_TEMPLATE.replace("__USER_DESKTOP__", USER_DESKTOP)

# 兼容旧代码
TOOLS = TOOL_EXECUTORS
has_tool_call = lambda text: bool(re.search(r'\{["\']tool["\']\s*:', text))


def _memory_base_paths():
    """记忆数据目录列表（平台感知路径，T3 检索式注入使用）。"""
    import os as _os
    from pathlib import Path as _Path
    dirs = []
    try:
        from src.cross_platform_engine import CrossPlatformStorage
        dirs.append(CrossPlatformStorage().base_path / "memories")
    except Exception:
        pass
    dirs.append(_Path.home() / ".meshctx" / "data" / "memories")
    # Windows 兼容：若在 WSL 也尝试 APPDATA 路径
    appdata = _os.environ.get("APPDATA")
    if appdata:
        dirs.append(_Path(appdata) / "meshctx" / "data" / "memories")
    return dirs


def _memory_item_score(item) -> float:
    """记忆条目注入排序分数 = importance × retention × schema_layer 加成。

    item 可为 dict（JSON 落盘）或 MemoryItem 对象。
    P3-4: semantic/core 层优先级高于 episodic（核心原则常驻，情景详情按需检索）
    —— 让三层图式化结构真正影响 token 分配。
    P2-1: retention 统一 10 底 + per-item FSRS stability（与 MemoryItem.current_retention
    同口径 R=10^(-t/S)，消除 e底/固定24h 与 FSRS 调度不一致）。
    """
    import time as _time
    get = item.get if isinstance(item, dict) else lambda k, d=None: getattr(item, k, d)
    imp = float(get("importance", 0.5) or 0.5)
    lr = get("last_reviewed") or get("last_accessed") or 0.0
    base = float(lr) if lr else float(get("created_at", 0.0) or 0.0)
    elapsed = max(0.0, _time.time() - base) if base else 0.0
    stability_h = float(get("stability", 24.0) or 24.0)  # FSRS stability（小时）
    s_seconds = max(1e-6, stability_h * 3600.0)
    # 统一 10 底（与 MemoryItem.current_retention 一致，审计点3）: R(t)=10^(-t/S)
    retention = max(0.05, min(1.0, 10.0 ** (-elapsed / s_seconds)))
    layer = get("schema_layer", None) or "episodic"
    layer_bonus = {"core": 1.25, "semantic": 1.15, "episodic": 1.0}.get(layer, 1.0)
    return imp * retention * layer_bonus


def _split_query_terms(text: str):
    """T3 轻量分词（零依赖）：英文按空格/非字母数字切，中文按 2-gram 滑动。

    解决中文 query 整串无法匹配的已知局限（2026-08-20 如实记录）。
    若环境装有 jieba 可替换为 jieba.lcut，但 2-gram 已够相关性粗筛。
    """
    import re as _re
    t = (text or "").lower()
    ascii_terms = [w for w in _re.split(r"[^a-z0-9]+", t) if len(w) >= 2]
    cjk_terms = []
    for seg in _re.findall(r"[\u4e00-\u9fff]+", t):
        if len(seg) >= 2:
            cjk_terms.extend(seg[i:i + 2] for i in range(len(seg) - 1))
    return ascii_terms + cjk_terms


def _maybe_review_memory_file(fp) -> None:
    """P2 复习闭环覆盖面（002 建议②）：prompt 注入命中的结构化记忆条目触发 FSRS 复习。

    与 _auto_review 同口径（grade=3 保守 + 防抖 1h）；回写失败不影响注入本身。
    """
    try:
        import json as _json2
        import time as _time2
        from pathlib import Path as _Path2
        from src.core.fsrs_scheduler import FSRSScheduler, MemoryCard
        fp = _Path2(fp)
        if not fp.exists():
            return
        m = _json2.loads(fp.read_text(encoding="utf-8"))
        now = _time2.time()
        last = float(m.get("last_reviewed", 0.0) or 0.0)
        if last and now - last < 3600.0:
            return  # 防抖 1h（与 _auto_review 一致）
        sched = FSRSScheduler()
        card = MemoryCard(
            item_id=m.get("id") or fp.stem,
            difficulty=float(m.get("difficulty", 5.0) or 5.0),
            stability=float(m.get("stability", 24.0) or 24.0) / 24.0,  # 小时→天
            interval_days=float(m.get("interval_days", 1.0) or 1.0),
            reviews=int(m.get("review_count", 0) or 0),
        )
        sched.review(card, grade=3)  # 保守 grade=3（与 record_recall 默认一致）
        m["stability"] = round(card.stability * 24.0, 2)  # 天→小时
        m["last_reviewed"] = now
        m["review_count"] = card.reviews
        m["difficulty"] = round(card.difficulty, 4)
        fp.write_text(_json2.dumps(m, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass  # 复习回写失败不影响记忆注入


def _collect_memory_entries(current_query: str = None, base_dirs: list = None, max_entries: int = 30):
    """T3: 检索式记忆注入 — 收集记忆条目并按 importance×retention 排序。

    - current_query 存在时：关键词相关性优先（top_k 检索，中英文均支持）
    - 否则：按 importance×retention 降序（最近/高频自然靠前）
    - 固定上限 max_entries，记忆段长度与记忆总量解耦。
    """
    import json as _json
    import time as _time
    from pathlib import Path as _Path
    rows = []  # list[dict]

    def _add(key, value, **meta):
        value = str(value or "").strip()
        if not value:
            return
        if any(r["value"] == value for r in rows):
            return
        rows.append({"key": key or "", "value": value, **meta})

    # 来源1: persistent_memory.json（纯文本 entries）
    mem_file = _Path.home() / ".meshctx" / "persistent_memory.json"
    if mem_file.exists():
        try:
            data = _json.loads(mem_file.read_text(encoding="utf-8"))
            if isinstance(data, dict) and data.get("entries"):
                for e in data["entries"]:
                    _add("", e, importance=0.5, last_reviewed=0.0, created_at=0.0)
        except Exception:
            pass

    # 来源2: memories/*.json（结构化记忆，带 importance/key）
    dirs = base_dirs if base_dirs is not None else _memory_base_paths()
    seen_dirs = set()
    for mem_dir in dirs:
        try:
            mem_dir = _Path(mem_dir)
            if not mem_dir.exists() or str(mem_dir) in seen_dirs:
                continue
            seen_dirs.add(str(mem_dir))
            for fp in sorted(mem_dir.glob("*.json")):
                try:
                    m = _json.loads(fp.read_text(encoding="utf-8"))
                    _add(m.get("key", ""), m.get("value") or m.get("content", ""),
                         importance=float(m.get("importance", 0.5) or 0.5),
                         last_reviewed=float(m.get("last_reviewed", 0.0) or 0.0),
                         created_at=float(m.get("created_at", 0.0) or 0.0),
                         stability=float(m.get("stability", 24.0) or 24.0),
                         schema_layer=m.get("schema_layer") or "episodic",
                         _src=str(fp))
                except Exception:
                    continue
        except Exception:
            continue

    # 排序：query 相关优先 → score 降序；无 query → score 降序
    if current_query:
        q = (current_query or "").lower()
        qterms = _split_query_terms(q)

        def _rel(r):
            hay = (r["key"] + " " + r["value"]).lower()
            hits = sum(1 for w in qterms if w in hay)
            # 完整 query 命中加权
            return hits + (2 if q and q in hay else 0)

        rows.sort(key=lambda r: (_rel(r), _memory_item_score(r)), reverse=True)
        # P2 复习闭环覆盖面（002 建议②）：注入 top-k 命中条目触发 FSRS 复习（防抖 1h）
        for r in rows[:max_entries]:
            src = r.get("_src")
            if src and _rel(r) > 0:
                _maybe_review_memory_file(src)
    else:
        rows.sort(key=lambda r: (_memory_item_score(r), r.get("created_at", 0.0)), reverse=True)

    return [r["value"] for r in rows[:max_entries]]


def build_system_prompt(project_dir: str = None, include_memory: bool = True,
                        current_query: str = None) -> str:
    """CLI 与 UI 共用的完整系统提示词 — 保证两端逐字一致。

    T1 前缀稳定化: system prompt 三段式——
      稳定段（SYSTEM_PROMPT + WSL 提示 + 项目上下文 + 工具定义）恒在最前、逐字节不变，
      记忆段按 key 稳定排序 + 固定上限（30 条），动态段最后。
      同任务多轮调用时 provider（DeepSeek/Anthropic/OpenAI）按前缀自动缓存命中。
    T3 检索式记忆注入: current_query 传入时按相关性 top_k 注入，记忆段长度恒定。
    """
    import platform as _platform
    from pathlib import Path as _Path
    parts = [SYSTEM_PROMPT]
    if "microsoft" in _platform.uname().release.lower():
        parts.append("\n⚠️ WSL 环境: C:\\ → /mnt/c/  D:\\ → /mnt/d/  可访问Windows文件。")
    if project_dir:
        proj = _Path(project_dir).expanduser()
        if proj.exists():
            for fname in ["AGENTS.md", "CODEBUDDY.md", "CLAUDE.md", ".cursorrules"]:
                fpath = proj / fname
                if fpath.exists():
                    parts.append(f"\n\n## 项目上下文 ({proj.name})\n{fpath.read_text(encoding='utf-8', errors='replace')[:3000]}")
                    break
    # 稳定段末尾: 工具定义（恒在记忆段之前 → 前缀稳定可缓存）
    parts.append("\n\n" + get_tools_prompt())
    # 记忆段: 检索式注入，按 importance×retention 排序，固定上限
    if include_memory:
        entries = _collect_memory_entries(current_query=current_query, max_entries=30)
        if entries:
            parts.append("\n\n## 🔒 持久化记忆（跨会话保留）\n\n")
            parts.append("\n".join(f"- {e}" for e in entries))
    return "".join(parts)


def trim_messages(messages: List[Dict], max_len: int = 40, keep: int = 30) -> List[Dict]:
    """公共消息清理：截断 + 保证 assistant(tool_calls)↔tool 配对完整。

    CLI 与 UI 统一后端共用，防止长对话/多轮工具调用后消息无限膨胀
    导致 token 超限。保留 system 在首位，剔除孤立的 tool 消息。
    """
    if len(messages) <= max_len:
        return messages
    system = messages[0] if messages and messages[0].get("role") == "system" else None
    recent = messages[-keep:] if keep else messages[-max_len:]
    # 去掉开头的孤立 tool 消息（缺少前置 assistant+tool_calls）
    while recent and recent[0].get("role") == "tool":
        recent.pop(0)
    # 收集所有 assistant 声明的 tool_call_id，过滤无匹配的孤立 tool 消息
    known_tool_call_ids = set()
    for m in recent:
        if m.get("role") == "assistant" and m.get("tool_calls"):
            for tc in m["tool_calls"]:
                known_tool_call_ids.add(tc.get("id", ""))
    filtered = [m for m in recent
                if not (m.get("role") == "tool"
                        and m.get("tool_call_id") not in known_tool_call_ids)]
    out = []
    if system:
        out.append(system)
    out.extend(filtered)
    return out

# ═══════════════════════════════════════════════════════════
# 统一 SYSTEM_PROMPT / TOOLS —— CLI 与 UI 共用同一份
# ═══════════════════════════════════════════════════════════

SYSTEM_PROMPT = """你是 meshctx AI 助手，运行在用户本机。

## 关于 meshctx

meshctx 是一个模块化 AI Agent 平台，开源版（当前运行）提供：
- **基础设施层（完整）**：沙箱执行、多通道通知（飞书/邮件/SMS）、认证、工作流编排、向量搜索、网页爬虫、日历引擎
- **AI 增强层（基础模式）**：脑启发路由、SDM 记忆、突破性记忆、OODA 循环框架
- **完整版能力**（需 meshctx-core 私有核心）：全局工作空间理论、多脑区竞争、自由能预测、JEPA 世界模型、多 Agent Swarm
- **你的角色**：你是 meshctx 的前端对话界面，直接帮助用户完成任务

当被问到 meshctx 自身架构时，诚实说明：开源版提供扎实的基础设施，高级 AI 能力（17脑区/意识点火等）在私有核心中。

## 最终答案格式

完成任务后，**最终交付必须输出"最终答案: "前缀**，后接简明、完整、具体的结果正文。
- 不要在最终答案中出现尖括号占位符（如 `<your answer>`、`<answer>`、`<output>`）——这些是待替换的模板，不是有效答案。
- 若任务要求"直接给出数字/一句话"，最终答案应只含结果，不做多余解释。

## 可用工具

| 工具 | 用途 |
|------|------|
| web_search | 搜索网页获取实时数据（⚠️ 用英文关键词，中文搜索效果差） |
| web_extract | 抓取指定网页的完整内容 |
| terminal   | 🔒 在本机执行 shell 命令（运行程序、安装软件、管理系统等） |
| browser_navigate | 抓取网页并提取可读文本（标题、链接、正文） |
| browser_snapshot | 获取已缓存页面的结构化内容 |
| read_file  | 读取本机文件 |
| write_file | 写入本机文件 |
| save_docx | 直接生成 Word(.docx) 报告（无需写脚本） |
| search_files | 搜索本机文件（按名称或内容） |
| remote_exec | 🔒 SSH 远程执行命令（需 host/user/password） |
| remote_read | 🔒 SSH 远程读取文件（需 host/user/password） |
| remote_write | 🔒 SSH 远程写入文件（需 host/user/password） |

## 搜索架构（多引擎自动回退）

web_search 使用5引擎链式回退：
1. **DDGS**（DuckDuckGo，默认）→ 失败自动跳到下一引擎
2. **Brave Search API**（需 BRAVE_API_KEY 环境变量）→ 免费2000次/月
3. **SearXNG**（需 SEARXNG_URL 环境变量，CloudCone自建）
4. **Google**（直接HTTP，Android移动端UA）
5. **Startpage**（Google代理，w-gl解析）

→ 一条路不通，下一条自动顶上。你只需正常调用 web_search，引擎切换对LLM透明。

## 本机目录速查（避免全盘扫描超时）

- Hermes 多 profile 目录: `~/.hermes/profiles/<profile>/`（如 quant / meshctx / codex / admin / bsc）
- meshctx 产品安装目录: `~/.meshctx/`；开发仓库: `~/meshctx-repo/`
- 用户主目录: `~`
- ⚠️ **禁止 `find /` 从根目录全盘扫描**——极慢，30 秒必然超时且无结果。找本机文件/目录用 `search_files` 工具，或直接 `ls` 已知路径（`~/.hermes`、`~/.meshctx`、`~` 等）

## 重要规则

- ⚠️ web_search 优先用英文关键词（多引擎对英文效果好）
- ⚠️ 如果连续2次返回「所有搜索引擎均失败」，立即停止搜索，改用 web_extract 抓取已知URL
- ⚠️ 本机 shell 命令 30 秒超时；若命令超时，换成更精确的路径/更小的范围，不要重复同一条慢命令
- 查询实时信息必须先调用 web_search
- 用户提供服务器信息（IP/用户名/密码）时，直接传入 remote_* 工具参数
- 读取/分析本机文件用 read_file
- 最终回复用中文，数据用表格呈现
- 被问到 meshctx 自身架构时，参考上方「关于 meshctx」诚实回答，不要搜索源码
- 生成 Word 报告/文档时：必须用 save_docx（直接生成 .docx）；禁止用 terminal/run_cmd 写 python 脚本生成报告（脚本硬编码路径会导致落到错误目录）。其他文件用 write_file。文件已存在时自动另存为新文件，不要覆盖旧文件；用绝对路径（用户桌面为 __USER_DESKTOP__）
- 用户要求「输出到桌面 / 生成 Word 报告 / 保存为文档」时：必须调用 save_docx 生成实际文件后才能结束；只输出文字总结视为未完成任务。"""

SYSTEM_PROMPT = SYSTEM_PROMPT.replace("__USER_DESKTOP__", USER_DESKTOP)

TOOLS = [
    {"type": "function", "function": {"name": "web_search", "description": "搜索网页获取实时信息（价格、新闻、天气等）。返回搜索结果摘要。", "parameters": {"type": "object", "properties": {"query": {"type": "string", "description": "搜索关键词"}}, "required": ["query"]}}},
    {"type": "function", "function": {"name": "web_extract", "description": "抓取指定 URL 的网页内容，返回纯文本。用于获取搜索结果的详细信息。", "parameters": {"type": "object", "properties": {"url": {"type": "string", "description": "要抓取的网页 URL"}}, "required": ["url"]}}},
    {"type": "function", "function": {"name": "read_file", "description": "读取本机文件。参数: path(文件路径), offset(起始行,默认1), limit(行数,默认200)", "parameters": {"type": "object", "properties": {"path": {"type": "string", "description": "文件路径"}, "offset": {"type": "integer", "description": "起始行号", "default": 1}, "limit": {"type": "integer", "description": "读取行数", "default": 200}}, "required": ["path"]}}},
    {"type": "function", "function": {"name": "write_file", "description": "写入本机文件。默认重名自动另存为新文件(文件名加 (1) 后缀)，不覆盖旧文件；只有明确修改已有文件时才传 if_exists='overwrite'。参数: path(文件路径), content(内容), if_exists(rename/overwrite)", "parameters": {"type": "object", "properties": {"path": {"type": "string", "description": "文件路径"}, "content": {"type": "string", "description": "要写入的内容"}, "if_exists": {"type": "string", "description": "目标已存在时的处理: rename=自动另存新文件(默认), overwrite=覆盖", "enum": ["overwrite", "rename"], "default": "rename"}}, "required": ["path", "content"]}}},
    {"type": "function", "function": {"name": "search_files", "description": "搜索本机文件（按名称或内容）。参数: pattern(搜索模式), dir(目录,默认HOME)", "parameters": {"type": "object", "properties": {"pattern": {"type": "string", "description": "搜索关键词或文件通配符"}, "dir": {"type": "string", "description": "搜索目录，默认用户 HOME"}}, "required": ["pattern"]}}},
    {"type": "function", "function": {"name": "terminal", "description": "🔒[需授权] 在本机执行 shell 命令。可以运行任何程序、脚本、安装软件、管理系统等。等同于用户在终端操作电脑。", "parameters": {"type": "object", "properties": {"cmd": {"type": "string", "description": "要执行的 shell 命令"}, "workdir": {"type": "string", "description": "工作目录(可选，默认当前目录)"}, "timeout": {"type": "integer", "description": "超时秒数(默认 60)", "default": 60}}, "required": ["cmd"]}}},
    {"type": "function", "function": {"name": "remote_read", "description": "🔒[需授权] 通过 SSH 读取远程服务器文件。用户提供服务器信息时请传入 host/user/password 参数。", "parameters": {"type": "object", "properties": {"path": {"type": "string", "description": "远程服务器文件路径"}, "host": {"type": "string", "description": "服务器地址"}, "user": {"type": "string", "description": "SSH 用户名"}, "password": {"type": "string", "description": "SSH 密码"}, "port": {"type": "integer", "description": "SSH 端口(默认 22)", "default": 22}}, "required": ["path", "host"]}}},
    {"type": "function", "function": {"name": "remote_write", "description": "🔒[需授权] 通过 SSH 写入远程服务器文件。用户提供服务器信息时请传入 host/user/password 参数。", "parameters": {"type": "object", "properties": {"path": {"type": "string", "description": "远程文件路径"}, "content": {"type": "string", "description": "要写入的内容"}, "host": {"type": "string", "description": "服务器地址"}, "user": {"type": "string", "description": "SSH 用户名"}, "password": {"type": "string", "description": "SSH 密码"}, "port": {"type": "integer", "description": "SSH 端口(默认 22)", "default": 22}}, "required": ["path", "content", "host"]}}},
    {"type": "function", "function": {"name": "remote_exec", "description": "🔒[需授权] 通过 SSH 在远程服务器执行命令。用户提供服务器信息时请传入 host/user/password 参数。", "parameters": {"type": "object", "properties": {"cmd": {"type": "string", "description": "要执行的 shell 命令"}, "host": {"type": "string", "description": "服务器地址"}, "user": {"type": "string", "description": "SSH 用户名"}, "password": {"type": "string", "description": "SSH 密码"}, "port": {"type": "integer", "description": "SSH 端口(默认 22)", "default": 22}}, "required": ["cmd", "host"]}}},
    {"type": "function", "function": {"name": "browser_navigate", "description": "抓取网页并提取可读文本（纯Python，无需Playwright）。参数: url", "parameters": {"type": "object", "properties": {"url": {"type": "string", "description": "网页URL"}}, "required": ["url"]}}},
    {"type": "function", "function": {"name": "browser_snapshot", "description": "获取当前已抓取页面的结构化内容（标题、链接、文本）。需先调用 browser_navigate", "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {"name": "save_memory", "description": "保存重要信息到持久化记忆，跨会话保留。用户告诉你重要信息（密码、配置、偏好等）时务必调用。", "parameters": {"type": "object", "properties": {"fact": {"type": "string", "description": "要记住的事实，一句话概括"}}, "required": ["fact"]}}},
    {"type": "function", "function": {"name": "save_docx", "description": "直接生成 Word(.docx) 报告/文档（服务器端 python-docx，无需写 python 脚本）。生成 Word 报告类文件时优先使用本工具。重名自动另存新文件(文件名加 (1) 后缀)，不覆盖旧文件。参数: path(绝对路径,以.docx结尾), title(报告标题), content(正文;每行一段,支持 #/##/### 标题和 - 列表)", "parameters": {"type": "object", "properties": {"path": {"type": "string", "description": "Word 文件绝对路径，如 {__USER_DESKTOP__}/报告.docx"}, "title": {"type": "string", "description": "报告标题"}, "content": {"type": "string", "description": "正文内容：每行一段；# 一级标题，## 二级，### 三级，- 列表项"}}, "required": ["path", "title", "content"]}}},
]


# 解析工具 schema 中的桌面占位符（跨平台：每台机器加载时注入真实桌面目录）
for _t in TOOLS:
    _fn = _t.get("function", {})
    _fn["description"] = _fn.get("description", "").replace("{__USER_DESKTOP__}", USER_DESKTOP)
    for _p in (_fn.get("parameters", {}) or {}).get("properties", {}).values():
        if isinstance(_p, dict) and "description" in _p:
            _p["description"] = _p["description"].replace("{__USER_DESKTOP__}", USER_DESKTOP)


# ═══════════════════════════════════════════════════════════
# 统一工具实现 —— CLI 与 UI 共用
# ═══════════════════════════════════════════════════════════

_BROWSER_CACHE = {}

def _web_extract(url: str) -> str:
    """抓取网页内容（支持 SOCKS5/HTTP 代理）"""
    import requests, re, os
    from src.config import load_config as _load_extract_config
    proxies = None
    try:
        cfg = _load_extract_config()
        proxy_url = cfg.get("search", {}).get("proxy", os.environ.get("MESHCTX_SEARCH_PROXY", ""))
    except Exception:
        proxy_url = os.environ.get("MESHCTX_SEARCH_PROXY", "")
    if proxy_url and proxy_url.strip():
        proxies = {"http": proxy_url, "https": proxy_url}
    try:
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        try:
            resp = requests.get(url, headers=headers, timeout=20, proxies=proxies, verify=False)
        except requests.exceptions.SSLError:
            import urllib3
            urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
            resp = requests.get(url, headers=headers, timeout=20, proxies=proxies, verify=False)
        html = resp.text
        text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text[:12000]
    except Exception as e:
        return f"抓取失败: {e}"


def _save_memory(fact: str) -> str:
    """保存事实到持久化记忆文件"""
    import json as _json_mem
    from pathlib import Path as _Path
    mem_file = _Path.home() / ".meshctx" / "persistent_memory.json"
    try:
        if mem_file.exists():
            data = _json_mem.loads(mem_file.read_text(encoding="utf-8"))
        else:
            data = {"entries": []}
        if fact.strip() not in data["entries"]:
            data["entries"].append(fact.strip())
            mem_file.parent.mkdir(parents=True, exist_ok=True)
            mem_file.write_text(_json_mem.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
            return f"✅ 已记住: {fact}"
        else:
            return f"已存在: {fact}"
    except Exception as e:
        return f"记忆保存失败: {e}"


def _terminal(cmd: str, workdir: str = "", timeout: int = 60) -> str:
    """本机执行 shell 命令（危险命令由 _run_cmd 内 CodeScanner 拦截）"""
    import os as _os
    if workdir:
        _cwd = _os.getcwd()
        try:
            _os.chdir(workdir)
        except Exception:
            pass
        result = _run_cmd(cmd)
        try:
            _os.chdir(_cwd)
        except Exception:
            pass
        return result
    return _run_cmd(cmd)


def _ssh_connect(host: str, user: str = "", password: str = "", port: int = 22):
    """建立 SSH 连接：参数 > 环境变量；paramiko → sshpass 回退"""
    import os as _os
    u = user or _os.environ.get("SERVER_USER", "root")
    pw = password or _os.environ.get("SERVER_PASS", "")
    if not pw:
        return None, host, u, "未提供密码（请传入 password 参数或设置 SERVER_PASS 环境变量）"
    try:
        import paramiko
        client = paramiko.SSHClient()
        client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        client.connect(host, port=port, username=u, password=pw, timeout=10, banner_timeout=5)
        return client, host, u, ""
    except Exception:
        pass
    return None, host, u, ""


def _ssh_exec_cmd(client, cmd: str, host: str = "", user: str = "", password: str = "", port: int = 22, timeout: int = 30) -> str:
    """在 SSH 客户端执行命令；client 为 None 时走 sshpass 回退"""
    import subprocess, shlex
    if client is not None:
        stdin, stdout, stderr = client.exec_command(cmd, timeout=timeout)
        out = stdout.read().decode(errors='replace').strip()
        err = stderr.read().decode(errors='replace').strip()
        return out or err
    import os as _os
    pw = password or _os.environ.get("SERVER_PASS", "")
    ssh_cmd = ["sshpass", "-p", pw, "ssh", "-o", "StrictHostKeyChecking=no", "-o", "ConnectTimeout=10", "-p", str(port), f"{user}@{host}", cmd]
    r = subprocess.run(ssh_cmd, shell=False, capture_output=True, text=True, timeout=timeout + 5)
    return (r.stdout + r.stderr).strip()


def _remote_exec(cmd: str, host: str = "", user: str = "", password: str = "", port: int = 22) -> str:
    """通过 SSH 执行远程命令"""
    client, h, u, err = _ssh_connect(host, user, password, port)
    if err:
        return f"远程执行失败: {err}"
    try:
        out = _ssh_exec_cmd(client, cmd, host=h, user=u, password=password, port=port)
        return out or "(无输出)"
    except Exception as e:
        return f"远程执行失败: {e}"
    finally:
        if client:
            client.close()


def _remote_read(path: str, host: str = "", user: str = "", password: str = "", port: int = 22) -> str:
    """通过 SSH 读取远程文件"""
    client, h, u, err = _ssh_connect(host, user, password, port)
    if err:
        return f"远程读取失败: {err}"
    try:
        out = _ssh_exec_cmd(client, f"cat '{path}'", host=h, user=u, password=password, port=port)
        lines = out.split('\n')
        return f"远程文件: {u}@{h}:{path} ({len(lines)} 行)\n" + '\n'.join(f"{i+1}|{l}" for i, l in enumerate(lines[:500]))
    except Exception as e:
        return f"远程读取失败: {e}"
    finally:
        if client:
            client.close()


def _remote_write(path: str, content: str, host: str = "", user: str = "", password: str = "", port: int = 22) -> str:
    """通过 SSH 写入远程文件"""
    client, h, u, err = _ssh_connect(host, user, password, port)
    if err:
        return f"远程写入失败: {err}"
    try:
        _ssh_exec_cmd(client, f"cat > '{path}' <<'MESHCTX_EOF'\n{content}\nMESHCTX_EOF\ncat '{path}' | wc -c", host=h, user=u, password=password, port=port)
        return f"已写入远程文件: {u}@{h}:{path} ({len(content)} 字符)"
    except Exception as e:
        return f"远程写入失败: {e}"
    finally:
        if client:
            client.close()


def _browser_navigate(url: str, cache: dict) -> str:
    """抓取网页并缓存（纯 HTTP 只读 fallback，与 UI 未授权路径一致）"""
    import re
    try:
        import requests
        from bs4 import BeautifulSoup
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, 'html.parser')
        title = soup.title.string.strip() if soup.title else "无标题"
        links = []
        for a in soup.find_all('a', href=True)[:20]:
            txt = a.get_text(strip=True)
            if txt:
                links.append(f"{txt} → {a['href']}")
        text = soup.get_text(separator=' ', strip=True)[:4000]
        cache[url] = {"title": title, "links": links, "text": text}
        return f"标题: {title}\n\n正文: {text}\n\n链接:\n" + "\n".join(links)
    except Exception as e:
        return f"抓取失败: {e}"


def _browser_snapshot(cache: dict) -> str:
    """返回已缓存页面内容"""
    if not cache:
        return "(缓存为空) 请先调用 browser_navigate"
    url, data = next(reversed(list(cache.items())))
    return f"当前页面: {url}\n标题: {data['title']}\n\n{data['text']}"
